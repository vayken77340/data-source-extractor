"""Give a specification template a real `Code` paragraph style.

The code blocks in a template — the JSON envelope of §6.4, the pseudocode, the rendered
keys — were built with *direct* formatting: a monospace font set run by run. That looks
right and is unusable, because changing it in Word means selecting every block by hand,
and the template is supposed to be yours to edit.

This adds a `Code` paragraph style, applies it to every paragraph that was formatted
monospace, and strips the direct formatting those paragraphs carried so the style
actually governs them. Afterwards, font, size, shading and spacing are one change in
Word's style pane and every block follows.

The style also fixes two things direct formatting could not:

- **keep-with-next**, so a 33-line envelope is not split across a page boundary.
- **a hanging indent**, so a key too long for the column wraps to a visibly deeper
  indent instead of continuing flush left where it reads as a new line.

Run it once per template. It is idempotent, and it works on a template you have already
customised — which is the point, since replacing your template would discard your edits.

    python tools/add_code_style.py config/specs/TEMPLATE.docx
    python tools/add_code_style.py config/specs/TEMPLATE.docx --font "Cascadia Mono"
"""

from __future__ import annotations

import argparse
import shutil
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

STYLE_ID = "Code"
STYLE_NAME = "Code"
DEFAULT_FONT = "Consolas"
DEFAULT_SIZE = 17  # half-points, so 8.5pt
SHADING = "F4F4F4"
BAR = "C8C8C8"
INDENT_LEFT = 624  # twips: where a wrapped line sits
INDENT_HANGING = 340  # twips: how far the first line sits back from it


def child(parent, tag: str, **attributes):
    """Replace `parent`'s `tag` child with a fresh one, so a rerun overwrites."""
    for existing in parent.findall(qn(tag)):
        parent.remove(existing)
    element = OxmlElement(tag)
    for name, value in attributes.items():
        element.set(qn(f"w:{name}"), value)
    parent.append(element)
    return element


def build_style(styles, font: str, size: int):
    """The `Code` style itself, replaced wholesale if it is already there."""
    root = styles.element
    for existing in root.findall(qn("w:style")):
        if existing.get(qn("w:styleId")) == STYLE_ID:
            root.remove(existing)

    style = OxmlElement("w:style")
    style.set(qn("w:type"), "paragraph")
    style.set(qn("w:styleId"), STYLE_ID)
    child(style, "w:name", val=STYLE_NAME)
    # Only base it on Normal when the template actually has one; a template built by
    # python-docx from a stripped source may not.
    if any(s.get(qn("w:styleId")) == "Normal" for s in root.findall(qn("w:style"))):
        child(style, "w:basedOn", val="Normal")
    child(style, "w:qFormat")

    paragraph = OxmlElement("w:pPr")
    child(paragraph, "w:keepNext")
    child(paragraph, "w:keepLines")
    borders = OxmlElement("w:pBdr")
    child(borders, "w:left", val="single", sz="18", space="6", color=BAR)
    paragraph.append(borders)
    child(paragraph, "w:shd", val="clear", color="auto", fill=SHADING)
    child(paragraph, "w:spacing", before="0", after="0", line="240", lineRule="auto")
    child(paragraph, "w:ind", left=str(INDENT_LEFT), hanging=str(INDENT_HANGING))
    style.append(paragraph)

    run = OxmlElement("w:rPr")
    child(run, "w:rFonts", ascii=font, hAnsi=font, cs=font)
    child(run, "w:sz", val=str(size))
    child(run, "w:szCs", val=str(size))
    style.append(run)

    root.append(style)
    return style


def is_monospaced(paragraph, font: str) -> bool:
    """A paragraph the template formatted as code, by the font its runs ask for."""
    for fonts in paragraph._p.iter(qn("w:rFonts")):
        if fonts.get(qn("w:ascii")) == font:
            return True
    return False


def apply(paragraph) -> None:
    """Put the paragraph under the style and remove what would override it."""
    paragraph.style = STYLE_ID
    properties = paragraph._p.find(qn("w:pPr"))
    if properties is not None:
        for tag in ("w:ind", "w:spacing"):
            for element in properties.findall(qn(tag)):
                properties.remove(element)
    # Direct run formatting wins over a style, so a font left on the run would make the
    # style's font a decoration. Size goes with it, for the same reason.
    for run in paragraph._p.iter(qn("w:rPr")):
        for tag in ("w:rFonts", "w:sz", "w:szCs"):
            for element in run.findall(qn(tag)):
                run.remove(element)


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description=__doc__.splitlines()[0])
    parser.add_argument("template", type=Path, help="the .docx to modify, in place")
    parser.add_argument("--font", default=DEFAULT_FONT, help=f"monospace font (default: {DEFAULT_FONT})")
    parser.add_argument("--size", type=float, default=DEFAULT_SIZE / 2, help="point size (default: 8.5)")
    parser.add_argument("--no-backup", action="store_true", help="do not write a .bak copy first")
    args = parser.parse_args(argv)

    if not args.template.is_file():
        raise SystemExit(f"no such file: {args.template}")
    if not args.no_backup:
        backup = args.template.with_suffix(args.template.suffix + ".bak")
        shutil.copy2(args.template, backup)
        print(f"backup  -> {backup}")

    document = Document(str(args.template))
    build_style(document.styles, args.font, int(round(args.size * 2)))

    touched = 0
    for paragraph in _every_paragraph(document):
        if is_monospaced(paragraph, args.font):
            apply(paragraph)
            touched += 1

    document.save(str(args.template))
    print(f"style   -> {STYLE_ID} ({args.font} {args.size}pt, shaded, keep-with-next, hanging indent)")
    print(f"applied -> {touched} paragraph(s) in {args.template}")
    if not touched:
        print(f"none matched: is the template's code font something other than {args.font!r}? see --font")
    return 0


def _every_paragraph(document):
    """Body paragraphs plus those inside tables, since callouts are one-cell tables."""
    yield from document.paragraphs
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs


if __name__ == "__main__":
    raise SystemExit(main())
